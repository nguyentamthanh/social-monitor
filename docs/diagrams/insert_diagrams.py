#!/usr/bin/env python3
"""Embed the 6 rendered diagram PNGs into the BA docx, right after the
content of their matching section (i.e. right before the next heading)."""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph
from docx.table import Table

HERE = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.normpath(os.path.join(HERE, '..', 'BA-Copyright-Checker.docx'))
PNG_DIR = os.path.join(HERE, 'diagrams', 'png')

PLACEMENTS = [
    # (heading prefix to anchor on, next-heading prefix that bounds the section,
    #  png filename, caption)
    ('7.1.', '7.2.', '01-tong-quan-kien-truc.png',
     'Hình 1. Tổng quan kiến trúc hệ thống — UI / API / Service / Data layer.'),
    ('FR-1:', 'FR-2:', '02-luong-xac-thuc.png',
     'Hình 2. Luồng xác thực đăng nhập/đăng ký (NextAuth Credentials + JWT).'),
    ('7.3.', '7.4.', '03-luong-quet-ban-quyen.png',
     'Hình 3. Luồng quét bản quyền — Quick Scan (URL/Find Copies) & Batch Scan.'),
    ('7.4.', '8.', '04-luong-url-check-find-copies.png',
     'Hình 4. Luồng Quick URL Check và Find Copies (YouTube).'),
    ('8.1.', '8.2.', '05-mo-hinh-du-lieu-erd.png',
     'Hình 5. Mô hình dữ liệu (ERD) — 8 bảng trong Neon Postgres.'),
    ('FR-10:', 'FR-11:', '06-luong-extension-api.png',
     'Hình 6. Luồng Extension API — xác thực bằng API key + kiểm tra video.'),
]


def iter_block_items(document):
    body = document.element.body
    for child in body.iterchildren():
        if child.tag.endswith('}p'):
            yield Paragraph(child, document)
        elif child.tag.endswith('}tbl'):
            yield Table(child, document)


def is_heading(block, prefix):
    if not isinstance(block, Paragraph):
        return False
    style = block.style.name if block.style else ''
    return style.startswith('Heading') and block.text.strip().startswith(prefix)


def block_element(block):
    return block._p if isinstance(block, Paragraph) else block._tbl


def insert_image_after(document, ref_element, image_path, caption, width_inches=6.3):
    img_p = document.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_p.add_run()
    run.add_picture(image_path, width=Inches(width_inches))

    cap_p = document.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_p.add_run(caption)
    cap_run.italic = True
    cap_run.font.size = Pt(9.5)
    cap_run.font.color.rgb = RGBColor(0x6d, 0x65, 0x8a)

    spacer_p = document.add_paragraph()

    # order after ref_element: img_p, cap_p, spacer_p
    ref_element.addnext(spacer_p._p)
    ref_element.addnext(cap_p._p)
    ref_element.addnext(img_p._p)


def main():
    document = Document(DOCX_PATH)

    for anchor_prefix, next_prefix, png_name, caption in PLACEMENTS:
        blocks = list(iter_block_items(document))
        anchor_idx = None
        for i, b in enumerate(blocks):
            if is_heading(b, anchor_prefix):
                anchor_idx = i
                break
        if anchor_idx is None:
            print(f"!! anchor not found: {anchor_prefix}")
            continue

        # find the next heading with next_prefix (search after anchor)
        end_idx = None
        for j in range(anchor_idx + 1, len(blocks)):
            if is_heading(blocks[j], next_prefix):
                end_idx = j
                break
        if end_idx is None:
            print(f"!! next-heading not found: {next_prefix} (after {anchor_prefix})")
            continue

        last_content_block = blocks[end_idx - 1]
        ref_element = block_element(last_content_block)
        image_path = os.path.join(PNG_DIR, png_name)
        insert_image_after(document, ref_element, image_path, caption)
        print(f"inserted {png_name} after section '{anchor_prefix}' (before '{next_prefix}')")

    document.save(DOCX_PATH)
    print("SAVED", DOCX_PATH)


if __name__ == '__main__':
    main()
